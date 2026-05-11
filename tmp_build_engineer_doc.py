from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

out='/Users/dan/Documents/airbnb-alerts/Cortex_Automation_Portfolio_ENGINEER.docx'

doc=Document()
styles=doc.styles
styles['Normal'].font.name='Calibri'
styles['Normal'].font.size=Pt(11)

p=doc.add_paragraph('CORTEX — INTERNAL BUILD BRIEF')
p.style='Title'
p.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT

doc.add_paragraph('Document: 8 Workflow Implementation Notes (Engineer Version)')
doc.add_paragraph('Audience: AI Engineer (Internal)')
doc.add_paragraph('Default Platform: Make (Integromat)')
doc.add_paragraph('Detail Level: Medium (step logic + key decisions, not exhaustive)')

intro=doc.add_paragraph()
intro.add_run('Purpose: ').bold=True
intro.add_run('This replaces client-facing copy with build-ready implementation notes. Each workflow includes scope, step logic, key decisions/gotchas, and a direct note from Omar.')

workflows=[
('01. Abandoned Cart Recovery (E-commerce)',[
'Build objective: Recover abandoned carts with timed multi-touch outreach and hard conversion checks before every send.',
'Core modules: Shopify webhook, Data Store (sequence state), Delay, Conditional checks, Email/SMS connectors, Logging sink (Airtable/Sheet/DB).',
'Step logic: 1) Capture abandonment event with cart value + customer contact. 2) Persist sequence state keyed by cart/session. 3) Wait 1h, re-check order status. 4) If not converted, send email #1. 5) Wait 24h and re-check; send email #2 if still open. 6) For high-value carts, add SMS branch at 24h or 72h. 7) On conversion event, close sequence and record recovered revenue.',
'Key decisions/gotchas: Deduplicate by cart/session ID; enforce quiet hours for SMS by timezone; do not send discount if SKU already discounted; add retry policy + dead-letter path for messaging API failures.',
"Omar note: Prioritise this first for e-commerce clients. Keep discount logic configurable per store (some clients want reminder-only with no incentive)."
]),
('02. AI Product Description Generator (E-commerce)',[
'Build objective: Generate brand-consistent, SEO-aware product copy in bulk and sync to draft listings for review.',
'Core modules: Source reader (Airtable/Google Sheets/PIM export), Prompt builder, LLM call, Validator, Shopify updater, Error queue.',
'Step logic: 1) Pull products flagged Needs Copy. 2) Assemble structured prompt from title, attributes, material, use-case, and brand tone. 3) Generate short + long description and optional SEO meta text. 4) Validate output length/forbidden claims. 5) Push to Shopify draft fields. 6) Mark row status and write token/cost usage.',
'Key decisions/gotchas: Guard against hallucinated specs; maintain blacklist for regulated claims; add human-review required flag for high-risk categories (health/children/electrical).',
"Omar note: Use reusable prompt templates by vertical (fashion, home, electronics). We need fast tuning without rebuilding the scenario."
]),
('03. AI-Powered Lead Scoring & Routing (SaaS/Tech)',[
'Build objective: Enrich inbound leads, score ICP fit + intent, and route to the correct path within minutes.',
'Core modules: HubSpot trigger, enrichment API (Clearbit/Apollo), scoring prompt/logic, router branches, Slack + CRM task creation.',
'Step logic: 1) Trigger on new lead in CRM. 2) Enrich company/person profile. 3) Compute deterministic pre-score (firm size, geo, segment) then LLM context score. 4) Combine into final score 1–10. 5) Route: hot => AE + Slack alert; warm => nurture; low => low-priority queue. 6) Persist reason codes for auditability.',
'Key decisions/gotchas: Keep model output structured JSON only; set timeout fallback to deterministic scoring if LLM fails; avoid routing loops when lead updates retrigger scenario.',
"Omar note: Accuracy of routing matters more than perfect AI text. Build transparent reason fields so sales trusts the score."
]),
('04. Churn Risk Detection & CSM Alert (SaaS/Tech)',[
'Build objective: Nightly churn scoring with actionable context for CSM follow-up.',
'Core modules: Scheduler, analytics fetcher (PostHog/Mixpanel), support signal pull, scoring function, Slack digest formatter, CRM task creator.',
'Step logic: 1) Nightly run pulls active accounts and last-30-day usage signals. 2) Calculate weighted risk score (logins, feature adoption, ticket velocity, seat utilisation). 3) Rank by ARR × risk. 4) Post top at-risk accounts to Slack digest. 5) Create owner-specific task in HubSpot with recommended next action.',
'Key decisions/gotchas: Define score versioning; track drift when product usage patterns change; protect against missing analytics events with null-safe defaults.',
"Omar note: Keep digest concise and ranked. CSMs will ignore it if it looks noisy. Top 10 accounts + why each is risky is enough."
]),
('05. AI Content Repurposing Pipeline (Marketing Agencies)',[
'Build objective: Convert one long-form asset into multiple channel-ready drafts with approval workflow.',
'Core modules: Content intake (Notion/GDoc/transcript), chunker, multi-branch LLM transforms, content calendar writer, approval trigger.',
'Step logic: 1) Intake long-form source. 2) Extract key points/quotes. 3) Fan out to channel-specific generation branches (LinkedIn, X thread, newsletter snippet, email hook, headlines). 4) Save outputs to Airtable/Notion with Pending Review status. 5) Notify approver in Slack; approval routes to scheduling queue.',
'Key decisions/gotchas: Enforce per-channel style constraints and character limits; include plagiarism/duplication checks against recent posts; maintain source citation link for editors.',
"Omar note: This one should feel assistant-like, not fully autonomous. Editors must stay in control before publish."
]),
('06. Client Reporting Automation (Marketing Agencies)',[
'Build objective: Auto-generate weekly performance reports from ad + CRM sources and send ready PDF deliverables.',
'Core modules: Scheduler, Meta/Google Ads fetchers, CRM metrics puller, transformation layer, Slides/Looker templater, PDF export + email.',
'Step logic: 1) Scheduled run starts Monday 7:00 AM local. 2) Pull prior-week ad and funnel metrics. 3) Normalize and map KPIs to template placeholders. 4) Generate slides/report PDF. 5) Send to account manager or directly to client list with intro block.',
'Key decisions/gotchas: Lock metric definitions across clients (avoid KPI drift); use date-window sanity checks; if any source fails, mark partial report and alert ops instead of sending bad data.',
"Omar note: Reliability is more important than fancy visuals. If data integrity check fails, hold send and alert team immediately."
]),
('07. Social Media Approval & Scheduling (Marketing Agencies)',[
'Build objective: Replace fragmented email approvals with structured approve/reject flow and direct scheduling.',
'Core modules: Draft detector (Airtable/Notion), Slack interactive approval, status updater, Buffer/native API scheduler, feedback router.',
'Step logic: 1) Detect new Draft item. 2) Create formatted approval message with content preview and metadata (platform/date/client). 3) Capture Approve / Request Edit response. 4) Approved branch schedules post and updates status. 5) Edit branch routes notes back to writer with context link.',
'Key decisions/gotchas: Ensure one-click actions are idempotent; prevent duplicate scheduling; enforce timezone correctness per client brand account.',
"Omar note: UX matters here. Make approval messages very clear so clients can approve in one click without confusion."
]),
('08. Inbound Lead Qualifier & Router (Marketing Agencies)',[
'Build objective: Respond to inbound leads instantly, qualify fit, and assign ownership with full context.',
'Core modules: Form webhook, enrichment step, qualification logic/LLM, immediate email responder, CRM create/update + assignment rules.',
'Step logic: 1) Trigger on form submit/booking event. 2) Enrich lead record (company, size, industry). 3) Score ICP fit + requested service alignment. 4) Send immediate personalised acknowledgement. 5) Route qualified leads to right AE queue with summary and recommended next step.',
'Key decisions/gotchas: Add anti-spam checks before enrichment/LLM spend; rate-limit by domain/IP; ensure SLA timestamps are stored for response-time reporting.',
"Omar note: Speed-to-lead is key. Keep end-to-end under 2 minutes for valid submissions."
]),
]

doc.add_paragraph('')
doc.add_paragraph('Contents', style='Heading 1')
for title,_ in workflows:
    doc.add_paragraph(title, style='List Number')

for title, bullets in workflows:
    doc.add_page_break()
    doc.add_paragraph(title, style='Heading 1')
    doc.add_paragraph('Implementation Notes', style='Heading 2')
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_paragraph('Build platform: Make (Integromat) default; n8n optional only when explicitly required.', style='Intense Quote')

doc.add_page_break()
doc.add_paragraph('Execution Guidance (Internal)', style='Heading 1')
for item in [
'Prioritisation suggestion: 01, 06, 08 first (fastest visible ROI).',
'Build standard: include error handler route, run logs, and retry policy in every scenario.',
'Documentation standard: handover should include scenario map, variables, webhook specs, and failure playbook.',
'Client customisation: keep thresholds/prompts/messages as data-driven config, not hardcoded logic.'
]:
    doc.add_paragraph(item, style='List Bullet')

doc.save(out)
print(out)
